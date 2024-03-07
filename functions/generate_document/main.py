import os
import json
from google.cloud import storage
from docx import Document
import io
from typing import Dict


def generate_document(event, context):
    bucket_name = event['bucket']
    minutes_file = event['name']

    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(minutes_file)

    minutes_json = blob.download_as_string().decode('utf-8')
    minutes = json.loads(minutes_json)

    document = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        document.add_heading(heading, level=1)
        document.add_paragraph(value)
        document.add_paragraph()

    document_bucket_name = os.environ.get('DOCUMENT_BUCKET_NAME')
    if not document_bucket_name:
        raise ValueError('Document bucket name not set')

    document_bucket = storage_client.bucket(document_bucket_name)
    document_blob = document_bucket.blob(f'{os.path.splitext(minutes_file)[0]}.docx')

    document_stream = io.BytesIO()
    document.save(document_stream)
    document_stream.seek(0)

    document_blob.upload_from_file(document_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
