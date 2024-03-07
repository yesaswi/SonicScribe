import os
from typing import Dict, Optional
from openai import OpenAI, OpenAIError
from docx import Document
from pydub import AudioSegment
from pydub.exceptions import CouldntDecodeError
from abc import ABC, abstractmethod
from pydantic import ValidationError, field_validator
from pydantic_settings import BaseSettings
import logging
import sys

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AppSettings(BaseSettings):
    openai_api_key: Optional[str] = os.environ.get("OPENAI_API_KEY")
    audio_model: str = "whisper-1"
    minutes_model: str = "gpt-4-0125-preview"
    audio_file_path: str
    transcription_file: str = "transcription.txt"
    minutes_file: str = "minutes.txt"
    output_file: str = "meeting_minutes.docx"

    class Config:
        env_file = ".env"
    
    @field_validator("audio_file_path")
    def validate_audio_file_path(cls, v):
        if v is None:
            raise ValueError("Audio file path is required.")
        if not os.path.isfile(v):
            raise FileNotFoundError(f"Audio file not found: {v}")
        if not v.lower().endswith((".wav", ".mp3", ".m4a")):
            raise ValueError(f"Invalid audio file format: {v}. Supported formats: .wav, .mp3, .m4a")
        return v

class AudioTranscriberError(Exception):
    pass

class MeetingMinutesExtractorError(Exception):
    pass

class MeetingMinutesWriterError(Exception):
    pass

class AudioTranscriber(ABC):
    @abstractmethod
    def transcribe_audio(self, audio_file_path: str, chunk_size_ms: int = 600_000) -> str:
        pass

class OpenAIAudioTranscriber(AudioTranscriber):
    def __init__(self, api_key: str, model: str):
        self.client = OpenAI(api_key=api_key)
        self.model = model

    def transcribe_audio(self, audio_file_path: str, chunk_size_ms: int = 600_000) -> str:
        try:
            with open(audio_file_path, 'rb') as audio_file:
                audio = AudioSegment.from_file(audio_file)
        except FileNotFoundError as e:
            raise AudioTranscriberError(f"Audio file not found: {audio_file_path}") from e
        except CouldntDecodeError as e:
            raise AudioTranscriberError(f"Could not decode audio file: {audio_file_path}") from e

        chunk_size = chunk_size_ms
        transcription = ""

        for i in range(0, len(audio), chunk_size):
            chunk = audio[i:i+chunk_size]
            chunk_file = "temp_chunk.wav"
            chunk.export(chunk_file, format="wav")

            with open(chunk_file, 'rb') as chunk_audio:
                try:
                    chunk_transcription = self.client.audio.transcriptions.create(
                        model=self.model,
                        file=chunk_audio,
                        response_format="json"
                    )
                    transcription += chunk_transcription.text
                except OpenAIError as e:
                    raise AudioTranscriberError(f"Error transcribing audio chunk: {e}") from e
                finally:
                    os.remove(chunk_file)

        return transcription

class MeetingMinutesExtractor(ABC):
    @abstractmethod
    def extract_minutes(self, transcription: str) -> Dict[str, str]:
        pass

class OpenAIMeetingMinutesExtractor(MeetingMinutesExtractor):
    def __init__(self, api_key: str, model: str):
        self.client = OpenAI(api_key=api_key)
        self.model = model

    def extract_minutes(self, transcription: str) -> Dict[str, str]:
        try:
            return {
                'abstract_summary': self._extract_abstract_summary(transcription),
                'key_points': self._extract_key_points(transcription),
                'action_items': self._extract_action_items(transcription),
                'sentiment': self._analyze_sentiment(transcription)
            }
        except OpenAIError as e:
            raise MeetingMinutesExtractorError(f"Error extracting meeting minutes: {e}") from e

    def _extract_abstract_summary(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points.",
            user_content=transcription
        )

    def _extract_key_points(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.",
            user_content=transcription
        )

    def _extract_action_items(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.",
            user_content=transcription
        )

    def _analyze_sentiment(self, transcription: str) -> str:
        return self._chat_completion(
            system_content="As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible.",
            user_content=transcription
        )

    def _chat_completion(self, system_content: str, user_content: str) -> str:
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                temperature=0,
                messages=[
                    {"role": "system", "content": system_content},
                    {"role": "user", "content": user_content}
                ]
            )
            return response.choices[0].message.content
        except OpenAIError as e:
            raise MeetingMinutesExtractorError(f"Error in chat completion: {e}") from e

class MeetingMinutesWriter:
    @staticmethod
    def save_as_docx(minutes: Dict[str, str], filename: str) -> None:
        try:
            doc = Document()
            for key, value in minutes.items():
                heading = ' '.join(word.capitalize() for word in key.split('_'))
                doc.add_heading(heading, level=1)
                doc.add_paragraph(value)
                doc.add_paragraph()
            doc.save(filename)
        except Exception as e:
            raise MeetingMinutesWriterError(f"Error saving meeting minutes: {e}") from e

def main():
    try:
        settings = AppSettings(audio_file_path=sys.argv[1] if len(sys.argv) > 1 else None)
        
        if settings.openai_api_key is None:
            raise ValueError("OpenAI API key is required.")
        
        if settings.audio_file_path is None:
            raise ValueError("Audio file path is required.")
    
    except FileNotFoundError as e:
        logger.error(f"Error: {str(e)}")
        return
    except ValidationError as e:
        logger.error(f"Invalid application settings: {e.errors()[0]['msg']}")
        return
    except ValueError as e:
        logger.error(str(e))
        return

    transcriber = OpenAIAudioTranscriber(settings.openai_api_key, settings.audio_model)
    minutes_extractor = OpenAIMeetingMinutesExtractor(settings.openai_api_key, settings.minutes_model)

    try:
        transcription = transcriber.transcribe_audio(settings.audio_file_path)
        with open(settings.transcription_file, "w") as file:
            file.write(transcription)
    except AudioTranscriberError as e:
        logger.error(str(e))
        return

    try:
        minutes = minutes_extractor.extract_minutes(transcription)
        with open(settings.minutes_file, "w") as file:
            file.write(str(minutes))
    except MeetingMinutesExtractorError as e:
        logger.error(str(e))
        return

    try:
        MeetingMinutesWriter.save_as_docx(minutes, settings.output_file)
    except MeetingMinutesWriterError as e:
        logger.error(str(e))
        return

if __name__ == "__main__":
    main()